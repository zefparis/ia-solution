"""
Module pour extraire et analyser le texte des fichiers PDF et Word (DOCX)
"""
import os
import tempfile
import logging
from datetime import datetime

# Configure le logger
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Importer les bibliothèques de traitement de documents
try:
    from docx import Document
    import pdfplumber
    from werkzeug.utils import secure_filename
    
    LIBRARIES_LOADED = True
except ImportError as e:
    logger.error(f"Erreur d'importation des bibliothèques de traitement de documents: {e}")
    LIBRARIES_LOADED = False

def extract_text_from_pdf(file_storage):
    """
    Extrait le texte d'un fichier PDF
    
    Args:
        file_storage: Flask FileStorage object
    
    Returns:
        dict: Résultat de l'extraction avec texte, métadonnées et statut
    """
    if not LIBRARIES_LOADED:
        return {
            "success": False,
            "error": "Les bibliothèques nécessaires ne sont pas installées."
        }
    
    try:
        # Créer un fichier temporaire
        temp_fd, temp_path = tempfile.mkstemp(suffix='.pdf')
        
        try:
            # Sauvegarde le fichier temporairement
            file_storage.save(temp_path)
            
            # Extrait le texte
            all_text = []
            metadata = {}
            
            with pdfplumber.open(temp_path) as pdf:
                # Extraire les métadonnées
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    metadata = {k.lower(): v for k, v in pdf.metadata.items() if v}
                
                # Extraire le texte de chaque page
                for page in pdf.pages:
                    text = page.extract_text(x_tolerance=3) or ""
                    all_text.append(text)
            
            # Nettoyer et joindre le texte
            full_text = "\n\n".join([t.strip() for t in all_text if t.strip()])
            
            return {
                "success": True,
                "text": full_text,
                "metadata": metadata,
                "document_type": "pdf"
            }
            
        finally:
            # Fermer et supprimer le fichier temporaire
            try:
                os.close(temp_fd)
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except Exception as e:
                logger.error(f"Erreur lors de la suppression du fichier temporaire: {e}")
    
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte PDF: {e}")
        import traceback
        logger.error(traceback.format_exc())
        
        return {
            "success": False,
            "error": str(e)
        }

def extract_text_from_docx(file_storage):
    """
    Extrait le texte d'un fichier Word (DOCX)
    
    Args:
        file_storage: Flask FileStorage object
    
    Returns:
        dict: Résultat de l'extraction avec texte, métadonnées et statut
    """
    if not LIBRARIES_LOADED:
        return {
            "success": False,
            "error": "Les bibliothèques nécessaires ne sont pas installées."
        }
    
    try:
        # Créer un fichier temporaire
        temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
        
        try:
            # Sauvegarde le fichier temporairement
            file_storage.save(temp_path)
            
            # Ouvrir le document
            doc = Document(temp_path)
            
            # Extraire les paragraphes
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extraire le texte des tableaux
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Joindre tous les textes
            all_text = paragraphs + tables_text
            full_text = "\n\n".join(all_text)
            
            # Extraire les métadonnées de base
            metadata = {}
            
            # Métadonnées des propriétés de documents (seulement certaines sont accessibles facilement)
            if hasattr(doc, 'core_properties'):
                if doc.core_properties.title:
                    metadata['title'] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata['author'] = doc.core_properties.author
                if doc.core_properties.created:
                    metadata['created'] = doc.core_properties.created.isoformat()
                if doc.core_properties.modified:
                    metadata['modified'] = doc.core_properties.modified.isoformat()
            
            return {
                "success": True,
                "text": full_text,
                "metadata": metadata,
                "document_type": "docx"
            }
            
        finally:
            # Fermer et supprimer le fichier temporaire
            try:
                os.close(temp_fd)
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
            except Exception as e:
                logger.error(f"Erreur lors de la suppression du fichier temporaire: {e}")
    
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        
        return {
            "success": False,
            "error": str(e)
        }

def is_allowed_file(filename, allowed_extensions=None):
    """
    Vérifie si un fichier a une extension autorisée
    
    Args:
        filename: Nom du fichier
        allowed_extensions: Liste d'extensions autorisées (par défaut: ['.pdf', '.docx'])
    
    Returns:
        bool: True si le fichier est autorisé, False sinon
    """
    if allowed_extensions is None:
        allowed_extensions = ['.pdf', '.docx']
    
    # Convertir en liste si c'est un tuple
    allowed_extensions = list(allowed_extensions)
    
    # Normaliser les extensions (s'assurer qu'elles commencent par un point)
    allowed_extensions = [ext if ext.startswith('.') else f'.{ext}' for ext in allowed_extensions]
    
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in [ext.lstrip('.').lower() for ext in allowed_extensions]

def extract_text_from_document(file_storage, filename=None):
    """
    Extrait le texte d'un document (PDF ou DOCX) en fonction de son extension
    
    Args:
        file_storage: Flask FileStorage object
        filename: Nom du fichier (optionnel, utilisera file_storage.filename si non fourni)
    
    Returns:
        dict: Résultat de l'extraction avec texte, métadonnées et statut
    """
    if not LIBRARIES_LOADED:
        return {
            "success": False,
            "error": "Les bibliothèques nécessaires ne sont pas installées."
        }
    
    if filename is None:
        filename = file_storage.filename
    
    if not filename:
        return {
            "success": False,
            "error": "Nom de fichier non spécifié"
        }
    
    # Vérifier l'extension
    if filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_storage)
    elif filename.lower().endswith('.docx'):
        return extract_text_from_docx(file_storage)
    else:
        return {
            "success": False,
            "error": f"Format de fichier non pris en charge: {filename}"
        }